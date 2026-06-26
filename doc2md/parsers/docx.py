"""python-docx parser — DOCX paragraphs + tables → markdown."""

from __future__ import annotations

import importlib

from .base import ExtractionResult, LibraryParser


class DocxParser(LibraryParser):
    name = "docx"
    module_name = "docx"
    extensions = (".docx",)

    def extract(self, source: str, *, ocr: bool = False) -> ExtractionResult:
        try:
            docx = importlib.import_module("docx")
            doc = docx.Document(source)
        except Exception as exc:
            return ExtractionResult(
                text="", parser=self.name, warnings=[f"python-docx failed: {exc}"]
            )

        out: list[str] = []
        # Paragraphs and tables interleave in document order via the body XML.
        from docx.oxml.ns import qn  # type: ignore

        body = doc.element.body
        para_map = {p._p: p for p in doc.paragraphs}
        table_map = {t._tbl: t for t in doc.tables}

        for child in body.iterchildren():
            if child in para_map:
                p = para_map[child]
                style = (p.style.name or "").lower() if p.style else ""
                txt = p.text.strip()
                if not txt:
                    continue
                if "heading 1" in style:
                    out.append(f"# {txt}")
                elif "heading 2" in style:
                    out.append(f"## {txt}")
                elif "heading 3" in style:
                    out.append(f"### {txt}")
                elif style.startswith("heading"):
                    out.append(f"#### {txt}")
                elif "list" in style:
                    out.append(f"- {txt}")
                else:
                    out.append(txt)
            elif child in table_map:
                out.append(_table_to_md(table_map[child]))

        return ExtractionResult(text="\n\n".join(out), parser=self.name)


def _table_to_md(table) -> str:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    md = ["| " + " | ".join(header) + " |",
          "| " + " | ".join("---" for _ in header) + " |"]
    for r in body:
        md.append("| " + " | ".join(r) + " |")
    return "\n".join(md)