#!/usr/bin/env python3
"""Generate APOC Section 11.3-11.9 as DOCX"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
heading = doc.add_heading('Section 11: APOC - Airport Operations Command Centre', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 11.3
doc.add_heading('11.3 WAISL AiOP Capability Summary', level=2)

doc.add_paragraph('The following table summarises the WAISL AiOP capability domains, what WAISL delivers for each domain, and the compliance status against the EAC RFP requirements.')

table1_data = [
    ['Capability Domain', 'What WAISL Delivers', 'EAC Compliance Status'],
    ['Real-Time Operations Intelligence', 'Unified operational picture across all SSH movements, updated in real time via event streaming from airport operational systems', 'Fully Compliant'],
    ['A-CDM Milestone Management', 'Full EUROCONTROL 16-milestone A-CDM tracking with cascade delay alerting and ATC integration (TOBT/TSAT)', 'Fully Compliant'],
    ['Passenger Load Forecasting', 'D0 to D+7 flight-level and terminal-zone passenger load forecasts using ML models', 'Fully Compliant'],
    ['What-If Scenario Simulation', 'Live-synchronised digital twin of SSH with scenario modelling before operational commitment', 'Fully Compliant'],
    ['Multi-Stakeholder Portals', 'Role-specific real-time dashboards for AOCC, airlines, ground handlers, ATC, and ECAA regulators', 'Fully Compliant'],
    ['Resource Demand-Capacity Balancing', 'Continuous demand vs. capacity computation at 15-minute resolution for all airport resources (D0-D+7)', 'Fully Compliant'],
    ['Baggage Operations Management', 'End-to-end bag tracking visibility, cutoff risk alerting, first/last bag SLA monitoring', 'Fully Compliant'],
    ['KPI & Reporting Framework', 'Operational KPI catalogue across five RFP-mandated domains; ten standard reports; three management horizons', 'Fully Compliant'],
    ['On-Premises Deployment', '100% on-premises at SSH on EAC-owned hardware; no cloud dependency; full EAC data sovereignty', 'Fully Compliant'],
    ['Security & Regulatory Compliance', 'EAC ICT policy, ECAA regulations, Egyptian NCSA NIS framework - defence-in-depth architecture', 'Fully Compliant'],
]

table1 = doc.add_table(rows=len(table1_data), cols=3)
table1.style = 'Table Grid'
for i, row_data in enumerate(table1_data):
    row = table1.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# 11.4
doc.add_heading('11.4 Data Integration Requirements', level=2)

doc.add_paragraph('WAISL AiOP requires the following operational data categories from airport operational systems. All integrations are delivered via the FAIRWAY Integration Bus, ensuring standardised, decoupled data exchange between source systems and the APOC platform.')

table2_data = [
    ['Source System', 'Data Categories Required', 'Interface Method', 'Refresh Rate'],
    ['DCS (LDCS)', 'Check-in completion events (pax count, bags, weight); Boarding scan events (gate, flight, timestamp); Flight load finalisation; No-show counts; BSM/BPM/PTM messages', 'FAIRWAY DX Connector / Type-B', 'Real-time (<5s)'],
    ['CUPPS', 'Counter allocation/deallocation events; Workstation session events; Boarding gate status; Check-in count per flight (aggregated)', 'FAIRWAY DX Connector', 'Near-real-time (<=30s)'],
    ['FIDS', 'Flight status updates (STD/STA, ETD/ETA, gate, stand, status code); Gate assignment/change events; Display confirmation events', 'FAIRWAY DX Connector', 'Real-time (<5s)'],
    ['AODB (INFOPAX)', 'Flight schedule (creation, update, cancellation); A-CDM milestones (EIBT, AIBT, TOBT, TSAT, TTOT, CTOT, AOBT, ATOT); Stand/gate allocations; Delay reason codes (IATA AHM 730); SSIM imports; ATFM slot assignments', 'FAIRWAY DX Connector', 'Real-time (<5s)'],
    ['BRS (BAGERA)', 'Baggage check-in/load events per flight; IATA Resolution 753 milestones (induction, sort, load, delivery); Belt assignments; First/last bag timestamps', 'FAIRWAY DX Connector', 'Real-time (<5s)'],
    ['RMS (INFOPAX-RMS)', 'Counter plan publications; Gate plan publications; Baggage belt assignments (flight, belt ID, carousel ID, allocation start, first/last bag times)', 'FAIRWAY DX Connector', 'Real-time (<10s)'],
    ['BHS', 'Bag induction events; Sortation decision events; Make-up unit loading progress (% inducted vs. TOBT); Screening status (cleared, held, re-screen); Loading completion events; Exception/missing bag alerts', 'FAIRWAY DX Connector / OPC-UA', 'Real-time (<5s)'],
]

table2 = doc.add_table(rows=len(table2_data), cols=4)
table2.style = 'Table Grid'
for i, row_data in enumerate(table2_data):
    row = table2.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# 11.5
doc.add_heading('11.5 Supported Data Exchange Standards', level=2)

doc.add_paragraph('WAISL AiOP supports the following industry-standard data exchange protocols and formats, ensuring compatibility with existing and future airport systems.')

table3_data = [
    ['Standard', 'Use Case'],
    ['FIXM 4.2', 'ATC flight data exchange (NANSC integration)'],
    ['IATA SSIM / SSIM2', 'Airline schedule data import'],
    ['AIDX', 'Flight information exchange'],
    ['IATA AHM 730', 'Delay attribution code framework'],
    ['IATA Resolution 753', 'Baggage tracking milestones'],
    ['BSM / BPM / PTM', 'Baggage service messages (DCS integration)'],
    ['SITA Type-B', 'Legacy airline DCS messaging (where applicable)'],
    ['REST / JSON', 'Primary API format for all modern integrations'],
    ['WebSocket', 'Real-time dashboard data feeds'],
    ['SFTP', 'Batch file exchange for schedules, reports'],
    ['METAR / TAF', 'Meteorological data (AFTN/REST)'],
]

table3 = doc.add_table(rows=len(table3_data), cols=2)
table3.style = 'Table Grid'
for i, row_data in enumerate(table3_data):
    row = table3.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# 11.6
doc.add_heading('11.6 APOC Services Available to EAC', level=2)

doc.add_paragraph('The following APOC capabilities are available for consumption by EAC stakeholders through standardised interfaces. These services form the basis of the operational intelligence delivered to AOCC operators, airlines, ground handlers, and regulatory bodies.')

table4_data = [
    ['APOC Capability', 'Description', 'Access Method'],
    ['Live Operations Dashboard', 'Real-time unified operational picture - all flights, resources, alerts, A-CDM milestones', 'WebSocket / REST API'],
    ['Passenger Load Forecasts', 'D0-D+7 passenger forecasts per flight and terminal zone with prediction confidence intervals', 'REST API (hourly update)'],
    ['Flight Operational Status', 'Full flight status including A-CDM milestones, TOBT, predicted EOBT, delay probability', 'REST API / WebSocket'],
    ['A-CDM Milestone Tracker', '16-milestone real-time tracking with target vs. actual and deviation alerts', 'REST API / WebSocket'],
    ['Resource Utilisation Status', 'Counter, gate, stand, belt utilisation at 15-minute resolution for D0-D+7', 'REST API'],
    ['Operational Alerts', 'Real-time alerts (delay cascade, counter saturation, baggage cutoff risk) with severity levels', 'WebSocket / Webhook'],
    ['Pre-Departure Sequence', 'TSAT-optimised departure queue for next 8 departures', 'REST API / WebSocket'],
    ['Standard Reports', 'Daily OTP, turnaround performance, baggage compliance, forecast accuracy - PDF and XLSX', 'Scheduled SFTP push / Web download'],
    ['KPI Dashboard Feed', 'Operational KPIs across punctuality, capacity utilisation, turnaround, and passenger experience', 'REST API (pre-computed)'],
    ['Regulatory Compliance Reports', 'IATA 753 baggage compliance, audit log, ECAA reporting data', 'REST API (scoped by role)'],
]

table4 = doc.add_table(rows=len(table4_data), cols=3)
table4.style = 'Table Grid'
for i, row_data in enumerate(table4_data):
    row = table4.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# 11.7
doc.add_heading('11.7 SLA Commitments', level=2)

doc.add_paragraph('WAISL commits to the following service level targets for the APOC platform. These SLAs are measured monthly and reported to EAC as part of the governance framework.')

table5_data = [
    ['SLA Area', 'Target', 'Notes'],
    ['Platform availability', '>= 99.9% monthly', 'Excludes planned maintenance and force majeure'],
    ['Operational event latency', '<= 30 seconds from source event to platform', 'Source-system outages excluded'],
    ['Dashboard refresh latency', '<= 2 seconds from platform event to display', 'Client browser performance excluded'],
    ['API response time (p95)', '< 500 ms', 'Batch jobs excluded'],
    ['Alert delivery', 'Within 30 seconds of trigger event', 'Pre-production environments excluded'],
    ['Data ingestion latency', '<= 5 seconds from source event to ingestion', 'Source-system outages excluded'],
    ['Forecast processing', 'Complete within 5 minutes per cycle', 'Initial model training excluded'],
    ['Passenger forecast accuracy D0', '<= 8% MAPE (day of operations)', 'SSH-specific calibration'],
    ['Passenger forecast accuracy D+7', '<= 25% MAPE', 'Confidence bands provided'],
    ['Delay attribution accuracy', '>= 95% within 4 hours of departure', 'IATA AHM 730 framework'],
    ['IATA 753 compliance tracking', '>= 99% of bags with full milestone data', '-'],
]

table5 = doc.add_table(rows=len(table5_data), cols=3)
table5.style = 'Table Grid'
for i, row_data in enumerate(table5_data):
    row = table5.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# 11.8
doc.add_heading('11.9 Deployment Model and Data Sovereignty', level=2)

doc.add_paragraph('WAISL AiOP is deployed 100% on-premises on EAC-owned infrastructure, ensuring full data sovereignty and operational independence. The following table summarises the deployment model commitments.')

table6_data = [
    ['Aspect', 'WAISL Commitment'],
    ['Deployment model', '100% on-premises at SSH'],
    ['Hardware ownership', 'All compute and storage run on EAC-owned infrastructure'],
    ['Data sovereignty', 'All EAC operational data remains on EAC-owned hardware at SSH - full regulatory compliance with Egyptian data residency requirements'],
    ['Vendor lock-in', 'EAC can continue operating the platform independently if the commercial relationship changes'],
    ['Future multi-site', 'Platform architecture is federation-ready for future extension to additional EAC airports without re-engineering'],
    ['Support model', 'WAISL is single point of accountability for APOC platform - integration, SLA, support, and all subcontractor coordination'],
]

table6 = doc.add_table(rows=len(table6_data), cols=2)
table6.style = 'Table Grid'
for i, row_data in enumerate(table6_data):
    row = table6.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# 11.9
doc.add_heading('11.9 APOC Compliance Matrix', level=2)

doc.add_paragraph('The full clause-by-clause APOC compliance matrix follows. All 53 requirements are met.')

# Create the 53-row compliance matrix
compliance_data = [
    ['Sr.', 'Requirement', 'Status', 'Remarks / WAISL Reference'],
    ['1', 'The Scope of the work includes design, construction, supply, integration, testing, commissioning, training documentation, maintenance and warranty of a single site - AOCC at Sharm El-Sheikh international Airport with an ICT and process architecture designed to support future evolution into a multisite control environment', 'Comply', 'Comply via WAISL AiOP. Platform architecture is federation-ready for multi-site extension.'],
    ['2', 'Implementation of the APOC platform to manage the airport operations plan and provide decision support for demand capacity balancing, pre-departure sequencing, VTT & delay analysis', 'Comply', 'Comply via WAISL AiOP. Standard capability with OPM and demand-capacity balancing engine.'],
    ['3', 'Integration with existing airport systems, including at minimum AODB, RMS, FIDS, slots/SSIM, A-CDM data, BHS', 'Comply', 'Comply via WAISL AiOP. FAIRWAY DX Connectors for all listed systems.'],
    ['4', 'Standardised API / Message bus layer (Service oriented / event driven architecture) to decouple resource, applications & external consumers, enabling future multi-site and multi airport data exchange', 'Comply', 'Comply via WAISL AiOP. REST API / WebSocket layer with OAuth 2.0 governance.'],
    ['5', 'Definition and configuration of APOC performance indicators, including KPIs and service levels for punctuality, predictability, capacity utilisation, turn around performance, and passenger experience aligned with recognised APOC performance catalogues', 'Comply', 'Comply via WAISL AiOP. KPI catalogue across five RFP-mandated domains.'],
    ['6', 'Implementation of dashboards, reports and data analytics tools to support daily management, post operations analysis, and continuous improvement', 'Comply', 'Comply via WAISL AiOP. Ten standard reports across three management horizons.'],
    ['7', "The proposed solution shall provide a real time, dynamic representation of the airport's operations with the following capabilities", 'Comply', 'Comply via WAISL AiOP. AOCC Master Operations Dashboard with six integrated views.'],
    ['8', 'Integrate seamlessly with existing systems to deliver actionable insights', 'Comply', 'Comply via WAISL AiOP. FAIRWAY Integration Bus decouples source systems from APOC.'],
    ['9', 'Support simulation capabilities to enhance operational efficiency', 'Comply', 'Comply via WAISL AiOP. What-If Scenario Simulation with seven scenario categories.'],
    ['10', 'Foster seamless collaboration amongst airport stakeholders and adhere to security and compliance standards.', 'Comply', 'Comply via WAISL AiOP. Multi-stakeholder portals with role-based access control.'],
    ['11', 'Implement advanced data analytic and machine learning techniques to develop a comprehensive data platform with dashboard.', 'Comply', 'Comply via WAISL AiOP. ML forecasting engine with automated retraining.'],
    ['12', 'Cover terminal and Concourse areas focusing on passenger touchpoint', 'Comply', 'Comply via WAISL AiOP. Passenger load forecasting per terminal zone.'],
    ['13', 'Provide a web visual representation of airport-wide operations to monitor, track, and simulate various operational scenarios.', 'Comply', 'Comply via WAISL AiOP. Web-based MFE dashboards with simulation engine.'],
    ['14', 'Proactively identify and address problems to ensure smooth and uninterrupted airport operations', 'Comply', 'Comply via WAISL AiOP. Real-time alerting with cascade delay prediction.'],
    ['15', 'Enable the airport to become a "data-driven" organization.', 'Comply', 'Comply via WAISL AiOP. KPI dashboard feed and standard reports.'],
    ['16', 'Drive improved operations and cost reduction.', 'Comply', 'Comply via WAISL AiOP. Resource optimisation and delay attribution.'],
    ['17', 'Support digital transformation of airport operations planning.', 'Comply', 'Comply via WAISL AiOP. Digital twin synchronised with live ops.'],
    ['18', 'System Integrations - Future capable to integrate seamlessly with disparate systems, incorporating IT, OT, IoT sensors, CUTE, CUSS, LDCS, structured and unstructured data.', 'Comply', 'Comply via WAISL AiOP. REST/SOAP/MQTT/OPC-UA/BACnet/Modbus support.'],
    ['19', 'Real-Time Monitoring - Real-time monitoring capabilities for passenger flow and operational activities (check-in, immigration, boarding gates), baggage movements on baggage reclaim belt, and alert/exception highlighting for service level breaches', 'Comply', 'Comply via WAISL AiOP. Dashboard refresh <=2s, alerts within 30s.'],
    ['20', 'What-If Simulation - Extensive what-if simulation capabilities allowing users to model and examine the implications of operational changes, including changes to operational model, resource restrictions, infrastructural capacity adjustments, and variations in arrival and departure rates.', 'Comply', 'Comply via WAISL AiOP. Seven scenario categories at go-live.'],
    ['21', 'Phase 2 Preparation - The platform shall be future capable to: support business rules and workflow configuration; provide generative AI capabilities on operational data; support self-service capabilities to manage the model and data; provide alerting and simulation model capabilities.', 'Comply', 'Comply via WAISL AiOP. Phase-2-ready architecture.'],
    ['22', 'The system shall forecast passenger load levels for both departing and arriving flights to enhance airport operational planning, resource allocation, and service efficiency through data driven passenger volume prediction for the D0 to D+7 period.', 'Comply', 'Comply via WAISL AiOP. D0-D+7 forecasts with <=8% MAPE (D0), <=25% MAPE (D+7).'],
    ['23', 'Predict the expected number of passengers for each departure and arrival flight.', 'Comply', 'Comply via WAISL AiOP. Flight-level and terminal-zone forecasts.'],
    ['24', 'Utilize historical passenger trends, flight schedules, seasonal patterns, and special event variations.', 'Comply', 'Comply via WAISL AiOP. ML models ingest historical trends, SSIM, seasonality.'],
    ['25', 'Incorporate real-time updates from airlines booking systems, check-in activity, and passenger manifests.', 'Comply', 'Comply via WAISL AiOP. Real-time refinement from CUPPS/DCS check-in data.'],
    ['26', 'The system shall integrate with: Airport Operational Database (AODB), Departure Control Systems (DCS), airline reservation and booking systems, and immigration and security throughput data (where permissible).', 'Comply', 'Comply via WAISL AiOP. FAIRWAY DX Connectors for AODB, DCS, CUPPS.'],
    ['27', 'Employ statistical and/or machine learning techniques to continuously improve prediction accuracy.', 'Comply', 'Comply via WAISL AiOP. Automated weekly model retraining.'],
    ['28', 'Provide defined accuracy targets for short-term (day of ops), midterm, and long-term forecast.', 'Comply', 'Comply via WAISL AiOP. <=8% MAPE (D0), <=25% MAPE (D+7) with confidence bands.'],
    ['29', 'Support automated model retraining with continuous performance monitoring.', 'Comply', 'Comply via WAISL AiOP. Weekly auto-retrain with MAPE tracking.'],
    ['30', 'Predictions shall support allocation of check-in counters, boarding gates, and security lanes; planning for immigrations staffing and passenger flow management; optimization of ground handling operations and terminal capacity usage.', 'Comply', 'Comply via WAISL AiOP. Resource demand-capacity balancing at 15-min resolution.'],
    ['31', 'Provide real-time dashboards showing forecasted vs. actual passenger loads', 'Comply', 'Comply via WAISL AiOP. AOCC Master Dashboard with pax load view.'],
    ['32', 'Offer configurable alerts for overcapacity or irregular passenger surges.', 'Comply', 'Comply via WAISL AiOP. Alert engine with severity classification.'],
    ['33', 'Enable report export in common formats such as PDF and Excel.', 'Comply', 'Comply via WAISL AiOP. Ten standard reports in PDF/XLSX.'],
    ['34', 'The solution shall fully support and integrate Airport Collaborative Decision Making (A-CDM) processes to enhance operational predictability, situational awareness, and coordination among all airport stakeholders.', 'Comply', 'Comply via WAISL AiOP. Full EUROCONTROL 16-milestone A-CDM framework.'],
    ['35', 'Targeted and actual operational timestamps.', 'Comply', 'Comply via WAISL AiOP. TOBT/TSAT target vs. actual tracking.'],
    ['36', 'Turnaround process data from airlines and ground handlers.', 'Comply', 'Comply via WAISL AiOP. Turnaround milestone ingestion from DCS/GHA.'],
    ['37', 'Resource availability and allocation information (stand, gate, crew, equipment)', 'Comply', 'Comply via WAISL AiOP. RMS integration for resource allocations.'],
    ['38', 'Adverse event notifications (delays, diversions, disruptions)', 'Comply', 'Comply via WAISL AiOP. Real-time alerting with cascade prediction.'],
    ['39', 'Pre-departure sequencing data from air traffic control', 'Comply', 'Comply via WAISL AiOP. ATC integration for CTOT/TSAT.'],
    ['40', 'Surface movement and flow data wherever available.', 'Comply', 'Comply via WAISL AiOP. A-SMGCS integration ready.'],
    ['41', 'The solution shall ingest, process, and analyse data streams using standard A-CDM exchange formats such as AIDX, FIXM, or other IATA/ICAO-aligned protocols.', 'Comply', 'Comply via WAISL AiOP. FIXM 4.2, AIDX, IATA SSIM support.'],
    ['42', 'The integrated solution shall use A-CDM data to enhance operational predictability.', 'Comply', 'Comply via WAISL AiOP. A-CDM milestone progression for delay prediction.'],
    ['43', 'Departure and arrival time predictions based on milestone progression.', 'Comply', 'Comply via WAISL AiOP. EOBT/TOBT prediction from milestone delta.'],
    ['44', 'Passenger load forecasts refined by real-time TOBT/TSAT dynamics.', 'Comply', 'Comply via WAISL AiOP. Pax forecast refinement on TOBT update.'],
    ['45', 'Baggage load predictions driven by updated check-in and turnaround milestones.', 'Comply', 'Comply via WAISL AiOP. BRS integration for bag load per flight.'],
    ['46', 'Turnaround progress monitoring with predictive alerts for deltas or risks.', 'Comply', 'Comply via WAISL AiOP. Turnaround Gantt with risk alerting.'],
    ['47', 'Forecasting resource utilization (gates, stands, check-in counters, security/immigration queues)', 'Comply', 'Comply via WAISL AiOP. Demand-capacity balancing D0-D+7.'],
    ['48', 'The system shall provide model retraining, anomaly detection, and continuous improvement of prediction accuracy.', 'Comply', 'Comply via WAISL AiOP. Weekly auto-retrain, anomaly detection engine.'],
    ['49', 'The A-CDM enabled prediction system shall provide actionable decision support for: AOCC, ATC, GHA, Airline Operations Centers, Baggage & Terminal Management teams', 'Comply', 'Comply via WAISL AiOP. Multi-stakeholder portals with role-based views.'],
    ['50', 'Real time deviation alerts when milestone progression deviates from expected patterns.', 'Comply', 'Comply via WAISL AiOP. Milestone deviation alerts within 30s.'],
    ['51', 'Early detection of delays and knock-on operational impacts', 'Comply', 'Comply via WAISL AiOP. Cascade delay prediction (90-min lookahead).'],
    ['52', 'Capacity and congestion forecasting including stands, gates, and terminal flow', 'Comply', 'Comply via WAISL AiOP. Resource utilisation forecasting at 15-min resolution.'],
    ['53', 'Predictive insights to support pre-departure sequencing and flow management.', 'Comply', 'Comply via WAISL AiOP. Pre-departure sequence queue (next 8 departures).'],
]

table7 = doc.add_table(rows=len(compliance_data), cols=4)
table7.style = 'Table Grid'
for i, row_data in enumerate(compliance_data):
    row = table7.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            if len(cell.paragraphs[0].runs) == 0:
                cell.paragraphs[0].add_run()
            cell.paragraphs[0].runs[0].bold = True
    # Set column widths for header row
    if i == 0:
        table7.columns[0].width = Cm(1.5)
        table7.columns[1].width = Cm(8)
        table7.columns[2].width = Cm(2.5)
        table7.columns[3].width = Cm(6)

# Save
output_path = '/Users/sujoymukherjee/code/doc2md/EAC_APOC_Section_11_3_to_11_9.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
